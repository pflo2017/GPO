from flask import Flask, render_template, request, jsonify, redirect, url_for, flash, current_app, g, session
from datetime import datetime, timedelta
import os
import re
from dotenv import load_dotenv
from faker import Faker
import random
import socket
import sys
import json
import logging
from werkzeug.utils import secure_filename
from flask_login import LoginManager, current_user, login_required
import uuid
from sqlalchemy import text

# Handle PyMuPDF import
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("Warning: PyMuPDF not available. PDF processing will be limited.")

# Handle docx import
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. DOCX processing will be limited.")

# Handle Google Generative AI import
try:
    import google.generativeai as genai
    # Check if the library has the necessary attributes or functions
    # Different versions of the library might have different APIs
    GENAI_AVAILABLE = True
    
    # Get API key from environment
    LLM_API_KEY = os.getenv('LLM_API_KEY')
    
    # Test if we can configure the API
    if LLM_API_KEY:
        try:
            if hasattr(genai, 'configure'):
                genai.configure(api_key=LLM_API_KEY)  # type: ignore
            # For newer versions that might use a different approach
            elif hasattr(genai, 'set_api_key'):
                genai.set_api_key(LLM_API_KEY)  # type: ignore
        except Exception as e:
            print(f"Warning: Failed to configure Google Generative AI: {str(e)}")
except ImportError:
    GENAI_AVAILABLE = False
    genai = None
    print("Warning: Google Generative AI library is not available")

# Import our modules
from auth import init_auth
from admin import init_admin
from database import db, init_database, Organization, User, Linguist, Project, ProjectDocument, AuditLog, Notification, LinguistProfile
from middleware import init_middleware
from api import api_bp
from forms import LinguistUploadForm, LinguistProfileForm, NewProjectRequestForm

# Load environment variables
load_dotenv()

# Create Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-key-for-testing')
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload size

# Ensure the upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize database
init_database(app)

fake = Faker()

# Configure logging for production
if app.config['DEBUG']:
    logging.basicConfig(level=logging.DEBUG)
else:
    logging.basicConfig(level=logging.INFO)

# Initialize authentication, admin, middleware, and API
init_auth(app)
init_admin(app)
init_middleware(app)
app.register_blueprint(api_bp)

# Set up multi-tenant context (now handled by middleware)
# The middleware handles organization context, user status checks, and audit logging

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/dashboard-debug')
@login_required
def dashboard_debug():
    """Debug version of dashboard to isolate rendering issues"""
    # Get projects for the current organization
    projects = Project.query.filter_by(organization_id=current_user.organization_id).order_by(Project.created_at.desc()).limit(10).all()
    linguists = LinguistProfile.query.filter_by(organization_id=current_user.organization_id).all()
    
    # Calculate statistics
    total_projects = len(projects)
    completed_analysis = Project.query.filter_by(
        organization_id=current_user.organization_id, 
        local_analysis_status='Analysis Complete'
    ).count()
    pending_analysis = Project.query.filter_by(
        organization_id=current_user.organization_id, 
        local_analysis_status='Pending Local Analysis'
    ).count()
    total_linguists = len(linguists)
    
    return render_template(
        'dashboard_debug.html',
        projects=projects,
        total_projects=total_projects,
        completed_analysis=completed_analysis,
        pending_analysis=pending_analysis,
        total_linguists=total_linguists
    )


@app.route('/dashboard')
@login_required
def dashboard():
    # Get projects for the current organization
    projects = Project.query.filter_by(organization_id=current_user.organization_id).order_by(Project.created_at.desc()).limit(10).all()
    linguists = LinguistProfile.query.filter_by(organization_id=current_user.organization_id).all()
    
    # Calculate statistics
    total_projects = len(projects)
    completed_analysis = Project.query.filter_by(
        organization_id=current_user.organization_id, 
        local_analysis_status='Analysis Complete'
    ).count()
    pending_analysis = Project.query.filter_by(
        organization_id=current_user.organization_id, 
        local_analysis_status='Pending Local Analysis'
    ).count()
    total_linguists = len(linguists)
    
    # Calculate AI risk distribution
    high_risk = Project.query.filter_by(
        organization_id=current_user.organization_id,
        ai_overall_risk_status='High'
    ).count()
    medium_risk = Project.query.filter_by(
        organization_id=current_user.organization_id,
        ai_overall_risk_status='Medium'
    ).count()
    low_risk = Project.query.filter_by(
        organization_id=current_user.organization_id,
        ai_overall_risk_status='Low'
    ).count()
    critical_risk = Project.query.filter_by(
        organization_id=current_user.organization_id,
        ai_overall_risk_status='Critical'
    ).count()
    
    return render_template(
        'dashboard.html',
        projects=projects,
        linguists=linguists,
        total_projects=total_projects,
        completed_analysis=completed_analysis,
        pending_analysis=pending_analysis,
        total_linguists=total_linguists,
        high_risk=high_risk,
        medium_risk=medium_risk,
        low_risk=low_risk,
        critical_risk=critical_risk
    )

@app.route('/analytics')
@login_required
def analytics():
    """Comprehensive analytics dashboard"""
    from sqlalchemy import func, desc
    from datetime import datetime, timedelta
    
    # Get basic analytics data
    projects = Project.query.filter_by(organization_id=current_user.organization_id).all()
    users = User.query.filter_by(organization_id=current_user.organization_id).all()
    
    # Calculate overview statistics
    total_projects = len(projects)
    active_projects = len([p for p in projects if p.status == 'In Progress'])
    completed_projects = len([p for p in projects if p.status == 'Completed'])
    total_users = len(users)
    
    # Calculate risk distribution
    risk_stats = db.session.query(
        Project.gpo_risk_status,
        func.count(Project.id)
    ).filter_by(organization_id=current_user.organization_id).group_by(Project.gpo_risk_status).all()
    
    risk_distribution = {status: count for status, count in risk_stats}
    
    # Calculate progress statistics
    total_words = sum(p.initial_word_count for p in projects)
    translated_words = sum(p.translated_words for p in projects)
    progress_percentage = (translated_words / total_words * 100) if total_words > 0 else 0
    
    # Get monthly trends (last 12 months)
    monthly_stats = db.session.query(
        func.date_trunc('month', Project.created_at).label('month'),
        func.count(Project.id)
    ).filter(
        Project.organization_id == current_user.organization_id,
        Project.created_at >= datetime.utcnow() - timedelta(days=365)
    ).group_by(func.date_trunc('month', Project.created_at)).order_by('month').all()
    
    # Get top clients
    client_stats = db.session.query(
        Project.client_name,
        func.count(Project.id).label('project_count')
    ).filter_by(organization_id=current_user.organization_id).group_by(
        Project.client_name
    ).order_by(desc('project_count')).limit(10).all()
    
    # Get recent projects
    recent_projects = Project.query.filter_by(
        organization_id=current_user.organization_id
    ).order_by(desc(Project.created_at)).limit(5).all()
    
    analytics_data = {
        'overview': {
            'total_projects': total_projects,
            'active_projects': active_projects,
            'completed_projects': completed_projects,
            'total_users': total_users
        },
        'risk_distribution': risk_distribution,
        'progress': {
            'total_words': total_words,
            'translated_words': translated_words,
            'progress_percentage': round(progress_percentage, 2)
        },
        'monthly_trends': [{
            'month': month.strftime('%Y-%m'),
            'count': count
        } for month, count in monthly_stats],
        'top_clients': [{
            'client_name': client,
            'project_count': count
        } for client, count in client_stats],
        'recent_projects': [{
            'id': p.id,
            'project_name': p.project_name,
            'status': p.status,
            'created_at': p.created_at.strftime('%Y-%m-%d')
        } for p in recent_projects]
    }
    
    return render_template('analytics.html', analytics=analytics_data)



@app.route('/project/<int:project_id>')
@login_required
def project_detail(project_id):
    project = Project.query.filter_by(id=project_id, organization_id=current_user.organization_id).first_or_404()
    return render_template('project_detail.html', project=project)

@app.route('/upload_document/<int:project_id>', methods=['POST'])
@login_required
def upload_document(project_id):
    project = Project.query.filter_by(id=project_id, organization_id=current_user.organization_id).first_or_404()
    
    if 'document' not in request.files:
        flash('No file part', 'error')
        return redirect(url_for('project_detail', project_id=project_id))
    
    file = request.files['document']
    if file.filename == '':
        flash('No selected file', 'error')
        return redirect(url_for('project_detail', project_id=project_id))
    
    if file and file.filename:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{project.id}_{filename}")
        file.save(file_path)
        
        # Count words in the document
        word_count = count_words(file_path)
        
        # Create project document
        document = ProjectDocument()
        document.project_id = project.id
        document.file_name = filename
        document.file_path = file_path
        document.file_type = os.path.splitext(filename)[1]
        document.word_count = word_count
        
        db.session.add(document)
        
        # Update project word count
        project.initial_word_count += word_count
        
        # Analyze the document and update risk status
        analyze_project(project, file_path)
        
        db.session.commit()
        
        flash('Document uploaded successfully', 'success')
    
    return redirect(url_for('project_detail', project_id=project_id))

@app.route('/assign_linguist/<int:project_id>', methods=['POST'])
@login_required
def assign_linguist(project_id):
    project = Project.query.filter_by(id=project_id, organization_id=current_user.organization_id).first_or_404()
    linguist_id = request.form.get('linguist_id')
    
    if linguist_id:
        linguist = Linguist.query.filter_by(id=linguist_id, organization_id=current_user.organization_id).first()
        if linguist:
            project.assigned_linguist_id = linguist.id
            db.session.commit()
            flash('Linguist assigned successfully', 'success')
        else:
            flash('Invalid linguist', 'error')
    
    return redirect(url_for('project_detail', project_id=project_id))

@app.route('/update_status/<int:project_id>', methods=['POST'])
@login_required
def update_status(project_id):
    project = Project.query.filter_by(id=project_id, organization_id=current_user.organization_id).first_or_404()
    status = request.form.get('status')
    
    if status:
        project.status = status
        db.session.commit()
        flash('Status updated successfully', 'success')
    
    return redirect(url_for('project_detail', project_id=project_id))

@app.route('/update_progress/<int:project_id>', methods=['POST'])
@login_required
def update_progress(project_id):
    project = Project.query.filter_by(id=project_id, organization_id=current_user.organization_id).first_or_404()
    translated_words = request.form.get('translated_words')
    
    if translated_words:
        try:
            translated_words = int(translated_words)
            if translated_words > project.initial_word_count:
                flash('Translated words cannot exceed initial word count', 'error')
            else:
                project.translated_words = translated_words
                
                # Update status based on progress
                if translated_words == 0:
                    project.status = 'Not Started'
                elif translated_words == project.initial_word_count:
                    project.status = 'Completed'
                else:
                    project.status = 'In Progress'
                
                db.session.commit()
                flash('Progress updated successfully', 'success')
        except ValueError:
            flash('Invalid word count', 'error')
    
    return redirect(url_for('project_detail', project_id=project_id))

# Utility functions
def count_words(file_path):
    """Count words in a document"""
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == '.pdf' and PYMUPDF_AVAILABLE:
            # Count words in PDF
            doc = fitz.open(file_path)
            word_count = 0
            for page in doc:
                # Extract text using the appropriate method based on PyMuPDF version
                text = ""
                # Try different methods to get text based on PyMuPDF version
                for method_name in ['get_text', 'getText', 'extractText']:
                    if hasattr(page, method_name):
                        try:
                            text = getattr(page, method_name)()
                            break
                        except:
                            continue
                
                words = text.split()
                word_count += len(words)
            return word_count
        
        elif (ext in ['.doc', '.docx']) and DOCX_AVAILABLE:
            # Count words in Word document
            doc = docx.Document(file_path)
            word_count = 0
            for para in doc.paragraphs:
                words = para.text.split()
                word_count += len(words)
            return word_count
        
        elif ext in ['.txt', '.md', '.html', '.xml']:
            # Count words in text file
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
                words = text.split()
                return len(words)
        
        else:
            # Default to a reasonable estimate
            return 1000
    
    except Exception as e:
        app.logger.error(f"Error counting words: {e}")
        return 1000

# Wrapper for Google Generative AI functionality
def analyze_with_genai(sample_text, project):
    """
    Analyze project text with Google's Generative AI
    """
    if not GENAI_AVAILABLE or not LLM_API_KEY:
        return {
            'risk_status': 'Unknown',
            'risk_reason': 'AI analysis not available',
            'recommendation': 'Please install the Google Generative AI library and provide an API key.'
        }
    
    try:
        prompt = f"""
        You are an AI assistant for a translation project management system.
        Analyze the following text sample from a translation project and assess the risk level.
        
        Project Details:
        - Language Pair: {project.language_pair}
        - Content Type: {project.content_type}
        - Initial Word Count: {project.initial_word_count}
        - Due Date: {project.due_date}
        
        Text Sample:
        {sample_text[:1000]}  # Limit sample size
        
        Provide a JSON response with the following fields:
        1. risk_status: Either "Low Risk", "Medium Risk", or "High Risk"
        2. risk_reason: A brief explanation of why this risk level was assigned
        3. recommendation: A recommendation for handling this project
        """
        
        # Try different ways to generate content based on available API
        response = None
        
        try:
            # For newer versions of the library
            if hasattr(genai, 'GenerativeModel'):
                model = genai.GenerativeModel('gemini-pro')  # type: ignore
                result = model.generate_content(prompt)  # type: ignore
                if hasattr(result, 'text'):
                    response = result.text  # type: ignore
                elif hasattr(result, 'parts') and len(result.parts) > 0:
                    response = result.parts[0].text  # type: ignore
            # For older versions of the library
            elif hasattr(genai, 'generate_text'):
                result = genai.generate_text(prompt=prompt)  # type: ignore
                if hasattr(result, 'result'):
                    response = result.result  # type: ignore
        except Exception as e:
            current_app.logger.error(f"Error generating AI content: {str(e)}")
            return {
                'risk_status': 'Unknown',
                'risk_reason': f'Error in AI analysis: {str(e)}',
                'recommendation': 'Please try again later or contact support.'
            }
        
        if not response:
            return {
                'risk_status': 'Unknown',
                'risk_reason': 'Failed to generate AI response',
                'recommendation': 'Please try again later or contact support.'
            }
            
        # Try to parse JSON from the response
        try:
            # Look for JSON in the response
            import re
            json_match = re.search(r'(\{.*\})', response.replace('\n', ' '), re.DOTALL)
            if json_match:
                json_str = json_match.group(1)
                result = json.loads(json_str)
            else:
                # If no JSON found, create a basic structure from the text
                result = {
                    'risk_status': 'Medium Risk',
                    'risk_reason': 'AI analysis completed but structured data not available',
                    'recommendation': response[:200] if response else 'No specific recommendation available'
                }
                
            # Validate and sanitize the result
            if 'risk_status' not in result or result['risk_status'] not in ['Low Risk', 'Medium Risk', 'High Risk', 'Unknown']:
                result['risk_status'] = 'Medium Risk'
            
            return result
        except Exception as e:
            current_app.logger.error(f"Error parsing AI response: {str(e)}")
            return {
                'risk_status': 'Medium Risk',
                'risk_reason': 'Error parsing AI response',
                'recommendation': 'Manual review recommended'
            }
    except Exception as e:
        current_app.logger.error(f"Error in AI analysis: {str(e)}")
        return {
            'risk_status': 'Unknown',
            'risk_reason': f'Error: {str(e)}',
            'recommendation': 'Please try again later.'
        }

def analyze_project(project, file_path):
    """Analyze a project and set risk status"""
    # Calculate days until due
    days_until_due = (project.due_date - datetime.now().date()).days
    
    # Calculate words per day needed
    words_per_day = project.initial_word_count / max(1, days_until_due) if days_until_due > 0 else float('inf')
    
    # Set risk status based on words per day
    if words_per_day > 3000:
        project.gpo_risk_status = 'High Risk'
        project.gpo_risk_reason = f'This project requires translating {int(words_per_day)} words per day, which exceeds normal capacity.'
        project.gpo_recommendation = 'Consider assigning multiple linguists or extending the deadline.'
    elif words_per_day > 2000:
        project.gpo_risk_status = 'Medium Risk'
        project.gpo_risk_reason = f'This project requires translating {int(words_per_day)} words per day, which is challenging.'
        project.gpo_recommendation = 'Monitor progress closely and be prepared to add resources if needed.'
    elif days_until_due < 0:
        project.gpo_risk_status = 'High Risk'
        project.gpo_risk_reason = 'This project is already past its due date.'
        project.gpo_recommendation = 'Prioritize completion and communicate delay to client.'
    else:
        project.gpo_risk_status = 'Low Risk'
        project.gpo_risk_reason = f'This project requires translating {int(words_per_day)} words per day, which is manageable.'
        project.gpo_recommendation = 'Proceed as planned with regular progress checks.'
    
    # If we have a LLM API key and Google Generative AI is available, use it for more advanced analysis
    if LLM_API_KEY and GENAI_AVAILABLE and genai:
        # Extract text from the document (limited sample)
        sample_text = extract_text_sample(file_path)
        
        # Use the wrapper function for Google Generative AI
        analysis = analyze_with_genai(sample_text, project)
        
        # Update project with AI analysis if available
        if analysis:
            project.gpo_risk_status = analysis.get('risk_status', project.gpo_risk_status)
            project.gpo_risk_reason = analysis.get('risk_reason', project.gpo_risk_reason)
            project.gpo_recommendation = analysis.get('recommendation', project.gpo_recommendation)

def extract_text_sample(file_path):
    """Extract text from a document file for analysis"""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # PDF extraction
        if file_ext == '.pdf':
            if PYMUPDF_AVAILABLE:
                try:
                    # Try to open the PDF with PyMuPDF
                    pdf = fitz.open(file_path)
                    text = ""
                    # Extract text from first few pages
                    max_pages = min(5, len(pdf))
                    for page_num in range(max_pages):
                        try:
                            # Try different methods depending on PyMuPDF version
                            page = pdf[page_num]
                            if hasattr(page, 'get_text'):
                                text += page.get_text()  # type: ignore
                            elif hasattr(page, 'getText'):
                                text += page.getText()  # type: ignore
                            else:
                                text += page.extractText()  # type: ignore
                        except Exception as e:
                            current_app.logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                    
                    # Limit text length
                    return text[:5000]
                except Exception as e:
                    current_app.logger.error(f"Error extracting text from PDF: {str(e)}")
                    return "Error extracting PDF text: " + str(e)
            else:
                return "PDF text extraction not available. Please install PyMuPDF."
        
        # DOCX extraction
        elif file_ext == '.docx':
            if DOCX_AVAILABLE:
                try:
                    doc = docx.Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return text[:5000]  # Limit text length
                except Exception as e:
                    current_app.logger.error(f"Error extracting text from DOCX: {str(e)}")
                    return "Error extracting DOCX text: " + str(e)
            else:
                return "DOCX text extraction not available. Please install python-docx."
        
        # TXT files
        elif file_ext == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    return f.read(5000)  # Read up to 5000 chars
            except Exception as e:
                current_app.logger.error(f"Error reading text file: {str(e)}")
                return "Error reading text file: " + str(e)
        
        # Unsupported format
        else:
            return f"Text extraction not supported for {file_ext} files."
    
    except Exception as e:
        current_app.logger.error(f"Error in extract_text_sample: {str(e)}")
        return "Error extracting text: " + str(e)

# Linguist Profile Management Routes
@app.route('/linguists')
@login_required
def linguists():
    """Display all linguist profiles for the current organization"""
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    
    # Build query with search filter
    query = LinguistProfile.query.filter_by(organization_id=current_user.organization_id)
    
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            db.or_(
                LinguistProfile.full_name.ilike(search_term),
                LinguistProfile.internal_id.ilike(search_term),
                LinguistProfile.specializations.ilike(search_term),
                LinguistProfile.source_languages.ilike(search_term),
                LinguistProfile.target_languages.ilike(search_term)
            )
        )
    
    # Paginate results
    linguists = query.order_by(LinguistProfile.full_name).paginate(
        page=page, per_page=20, error_out=False
    )
    
    return render_template('linguists.html', linguists=linguists, search=search)

@app.route('/linguists/upload', methods=['GET', 'POST'])
@login_required
def upload_linguists():
    """Handle linguist profile upload via CSV/Excel file"""
    form = LinguistUploadForm()
    
    if form.validate_on_submit():
        file = form.file.data
        filename = file.filename.lower()
        
        try:
            import pandas as pd
            import io
            
            # Read the file based on its extension
            if filename.endswith('.csv'):
                df = pd.read_csv(file)
            elif filename.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file)
            else:
                flash('Unsupported file format. Please upload CSV or Excel files.', 'error')
                return redirect(url_for('upload_linguists'))
            
            # Expected columns (case-insensitive)
            expected_columns = [
                'internal_id', 'full_name', 'email', 'specializations',
                'source_languages', 'target_languages', 'quality_rating',
                'general_capacity_words_per_day', 'status'
            ]
            
            # Normalize column names to lowercase
            df.columns = df.columns.str.lower().str.strip()
            
            # Check for required columns
            missing_columns = [col for col in ['internal_id', 'full_name', 'source_languages', 'target_languages'] 
                             if col not in df.columns]
            if missing_columns:
                flash(f'Missing required columns: {", ".join(missing_columns)}', 'error')
                return redirect(url_for('upload_linguists'))
            
            if df.empty:
                flash('The uploaded file is empty.', 'warning')
                return redirect(url_for('upload_linguists'))
            
            uploaded_count = 0
            updated_count = 0
            error_count = 0
            
            for index, row in df.iterrows():
                try:
                    # Clean and validate data using explicit pandas methods
                    internal_id_raw = row.get('internal_id', '')
                    full_name_raw = row.get('full_name', '')
                    email_raw = row.get('email')
                    specializations_raw = row.get('specializations')
                    source_languages_raw = row.get('source_languages', '')
                    target_languages_raw = row.get('target_languages', '')
                    quality_rating_raw = row.get('quality_rating')
                    capacity_raw = row.get('general_capacity_words_per_day')
                    status_raw = row.get('status', 'Active')
                    
                    # Convert to strings safely using explicit checks
                    internal_id = str(internal_id_raw).strip() if internal_id_raw is not None and str(internal_id_raw) != 'nan' else ""
                    full_name = str(full_name_raw).strip() if full_name_raw is not None and str(full_name_raw) != 'nan' else ""
                    email = str(email_raw).strip() if email_raw is not None and str(email_raw) != 'nan' else None
                    specializations = str(specializations_raw).strip() if specializations_raw is not None and str(specializations_raw) != 'nan' else None
                    source_languages = str(source_languages_raw).strip() if source_languages_raw is not None and str(source_languages_raw) != 'nan' else ""
                    target_languages = str(target_languages_raw).strip() if target_languages_raw is not None and str(target_languages_raw) != 'nan' else ""
                    quality_rating = str(quality_rating_raw).strip() if quality_rating_raw is not None and str(quality_rating_raw) != 'nan' else None
                    capacity = capacity_raw if capacity_raw is not None and str(capacity_raw) != 'nan' else None
                    status = str(status_raw).strip() if status_raw is not None and str(status_raw) != 'nan' else "Active"
                    
                    # Validate required fields
                    if (internal_id == "" or full_name == "" or source_languages == "" or target_languages == ""):
                        error_count += 1
                        continue
                    
                    # Check if linguist already exists
                    existing_linguist = LinguistProfile.query.filter_by(
                        organization_id=current_user.organization_id,
                        internal_id=internal_id
                    ).first()
                    
                    if existing_linguist:
                        # Update existing linguist
                        existing_linguist.full_name = full_name
                        existing_linguist.email = email
                        existing_linguist.specializations = specializations
                        existing_linguist.source_languages = source_languages
                        existing_linguist.target_languages = target_languages
                        existing_linguist.quality_rating = quality_rating
                        existing_linguist.general_capacity_words_per_day = capacity
                        existing_linguist.status = status
                        existing_linguist.updated_at = datetime.utcnow()
                        updated_count += 1
                    else:
                        # Create new linguist
                        new_linguist = LinguistProfile()
                        new_linguist.organization_id = current_user.organization_id
                        new_linguist.internal_id = internal_id
                        new_linguist.full_name = full_name
                        new_linguist.email = email
                        new_linguist.specializations = specializations
                        new_linguist.source_languages = source_languages
                        new_linguist.target_languages = target_languages
                        new_linguist.quality_rating = quality_rating
                        new_linguist.general_capacity_words_per_day = capacity
                        new_linguist.status = status
                        db.session.add(new_linguist)
                        uploaded_count += 1
                
                except Exception as e:
                    error_count += 1
                    current_app.logger.error(f"Error processing row {index}: {str(e)}")
            
            # Commit all changes
            db.session.commit()
            
            # Show results
            if uploaded_count > 0 or updated_count > 0:
                message = f"Successfully processed {uploaded_count} new linguists and updated {updated_count} existing linguists."
                if error_count > 0:
                    message += f" {error_count} rows had errors and were skipped."
                flash(message, 'success')
            else:
                flash('No linguists were processed. Please check your file format.', 'warning')
            
            return redirect(url_for('linguists'))
            
        except Exception as e:
            flash(f'Error processing file: {str(e)}', 'error')
            current_app.logger.error(f"File upload error: {str(e)}")
            return redirect(url_for('upload_linguists'))
    
    return render_template('upload_linguists.html', form=form)

@app.route('/linguists/create', methods=['GET', 'POST'])
@login_required
def create_linguist():
    """Create a new linguist profile manually"""
    form = LinguistProfileForm()
    
    if form.validate_on_submit():
        # Check if internal_id already exists for this organization
        existing = LinguistProfile.query.filter_by(
            organization_id=current_user.organization_id,
            internal_id=form.internal_id.data
        ).first()
        
        if existing:
            flash('A linguist with this Internal ID already exists.', 'error')
            return render_template('create_linguist.html', form=form)
        
        # Create new linguist
        new_linguist = LinguistProfile()
        new_linguist.organization_id = current_user.organization_id
        new_linguist.internal_id = form.internal_id.data
        new_linguist.full_name = form.full_name.data
        new_linguist.email = form.email.data
        new_linguist.specializations = form.specializations.data
        new_linguist.source_languages = form.source_languages.data
        new_linguist.target_languages = form.target_languages.data
        new_linguist.quality_rating = form.quality_rating.data
        new_linguist.general_capacity_words_per_day = form.general_capacity_words_per_day.data
        new_linguist.status = form.status.data
        
        db.session.add(new_linguist)
        db.session.commit()
        
        flash('Linguist profile created successfully!', 'success')
        return redirect(url_for('linguists'))
    
    return render_template('create_linguist.html', form=form)

@app.route('/linguists/<linguist_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_linguist(linguist_id):
    """Edit an existing linguist profile"""
    linguist = LinguistProfile.query.filter_by(
        id=linguist_id,
        organization_id=current_user.organization_id
    ).first_or_404()
    
    form = LinguistProfileForm(obj=linguist)
    
    if form.validate_on_submit():
        # Check if internal_id already exists for another linguist in this organization
        existing = LinguistProfile.query.filter(
            LinguistProfile.organization_id == current_user.organization_id,
            LinguistProfile.internal_id == form.internal_id.data,
            LinguistProfile.id != linguist_id
        ).first()
        
        if existing:
            flash('A linguist with this Internal ID already exists.', 'error')
            return render_template('edit_linguist.html', form=form, linguist=linguist)
        
        # Update linguist
        form.populate_obj(linguist)
        linguist.updated_at = datetime.utcnow()
        db.session.commit()
        
        flash('Linguist profile updated successfully!', 'success')
        return redirect(url_for('linguists'))
    
    return render_template('edit_linguist.html', form=form, linguist=linguist)

@app.route('/linguists/<linguist_id>/delete', methods=['POST'])
@login_required
def delete_linguist(linguist_id):
    """Delete a linguist profile"""
    linguist = LinguistProfile.query.filter_by(
        id=linguist_id,
        organization_id=current_user.organization_id
    ).first_or_404()
    
    db.session.delete(linguist)
    db.session.commit()
    
    flash('Linguist profile deleted successfully!', 'success')
    return redirect(url_for('linguists'))

# Project Analysis Request Routes
@app.route('/new_project_request', methods=['GET', 'POST'])
@login_required
def new_project_request():
    """Initiate a new project analysis request"""
    form = NewProjectRequestForm()
    
    # Populate linguist choices for the current organization
    linguists = LinguistProfile.query.filter_by(
        organization_id=current_user.organization_id,
        status='Active'
    ).order_by(LinguistProfile.full_name).all()
    
    # For SelectField choices assignment:
    choices = [('', 'Select a linguist (optional)')]
    choices.extend([
        (linguist.id, f"{linguist.full_name} ({linguist.internal_id}) - {linguist.source_languages} → {linguist.target_languages}")
        for linguist in linguists
    ])
    form.selected_linguist_id_for_planning.choices = list(choices)
    
    if form.validate_on_submit():
        # Create new project
        new_project = Project()
        new_project.client_name = form.client_name.data
        new_project.project_name = form.project_name.data
        new_project.source_lang = form.source_lang.data
        new_project.target_lang = form.target_lang.data
        new_project.content_type = form.content_type_selection.data
        new_project.desired_deadline = form.desired_deadline.data
        new_project.selected_linguist_id_for_planning = form.selected_linguist_id_for_planning.data if form.selected_linguist_id_for_planning.data else None
        new_project.organization_id = current_user.organization_id
        new_project.created_by = current_user.id
        new_project.local_analysis_status = 'Pending Local Analysis'
        
        db.session.add(new_project)
        db.session.commit()
        
        # Generate success message with instructions
        flash(f'Project analysis request created successfully! Local Analysis Request ID: {new_project.id}. Please use your GPO Local Brain to analyze the document with this ID.', 'success')
        return redirect(url_for('dashboard'))
    
    return render_template('new_project_request.html', form=form)

@app.route('/projects/<project_id>')
@login_required
def project_details(project_id):
    """View project details and AI blueprint"""
    project = Project.query.filter_by(
        id=project_id,
        organization_id=current_user.organization_id
    ).first_or_404()
    
    return render_template('project_details.html', project=project)

@app.route('/api/local-analysis-results', methods=['POST'])
def local_analysis_results():
    """API endpoint for receiving analysis results from Local Brain"""
    try:
        # Get API key from request headers
        api_key = request.headers.get('X-API-Key')
        if not api_key:
            return jsonify({'error': 'API key required'}), 401
        
        # Validate API key
        organization = Organization.query.filter_by(api_key=api_key).first()
        if not organization:
            return jsonify({'error': 'Invalid API key'}), 401
        
        # Parse JSON payload
        if not request.data or request.data.strip() == b'':
            return jsonify({'error': 'Invalid JSON payload'}), 400
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'error': 'Invalid JSON payload'}), 400
        
        # Extract required fields
        local_analysis_request_id = data.get('local_analysis_request_id')
        if not local_analysis_request_id:
            return jsonify({'error': 'local_analysis_request_id required'}), 400
        
        # Find the project
        project = Project.query.filter_by(
            id=local_analysis_request_id,
            organization_id=organization.id
        ).first()
        
        if not project:
            return jsonify({'error': 'Project not found'}), 404
        
        # Update project with AI analysis results
        project.ai_overall_risk_status = data.get('ai_overall_risk_status')
        project.ai_risk_reason = data.get('ai_risk_reason')
        project.ai_document_complexity = data.get('ai_document_complexity')
        project.ai_key_challenges = data.get('ai_key_challenges')
        project.ai_sensitive_data_alert_summary = data.get('ai_sensitive_data_alert_summary')
        project.ai_recommended_linguist_profile_text = data.get('ai_recommended_linguist_profile_text')
        project.ai_optimal_team_size = data.get('ai_optimal_team_size')
        project.ai_deadline_fit_assessment = data.get('ai_deadline_fit_assessment')
        project.ai_strategic_recommendations = data.get('ai_strategic_recommendations')
        project.local_analysis_status = data.get('local_analysis_status', 'Analysis Complete')
        
        # Parse timestamp if provided
        if data.get('ai_analysis_timestamp'):
            try:
                project.ai_analysis_timestamp = datetime.fromisoformat(data['ai_analysis_timestamp'].replace('Z', '+00:00'))
            except:
                project.ai_analysis_timestamp = datetime.utcnow()
        else:
            project.ai_analysis_timestamp = datetime.utcnow()
        
        project.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Analysis results updated successfully',
            'project_id': project.id
        })
        
    except Exception as e:
        current_app.logger.error(f"Error processing local analysis results: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500

# Error handlers
@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

@app.errorhandler(500)
def server_error(e):
    return render_template('500.html'), 500

@app.context_processor
def inject_now():
    from datetime import datetime
    return {'now': datetime.utcnow()}

if __name__ == '__main__':
    print('🚀 Starting GPO Application...')
    print('⏳ Checking database connection...')
    try:
        with app.app_context():
            db.engine.connect()
            print('🗄️  Database connection successful!')
    except Exception as e:
        print(f'❌ Database connection error: {e}')
    
    print('🔥 Starting Flask server...')
    port = 5001  # Use a different port to avoid conflicts with AirPlay on macOS
    print(f'🌐 Server will be available at http://localhost:{port}')
    app.run(host='0.0.0.0', port=port, debug=True) 